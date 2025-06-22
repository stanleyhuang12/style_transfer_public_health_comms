import requests
from bs4 import BeautifulSoup 
from urllib.parse import urlparse, urljoin
import nltk
import pytesseract
import io 
import numpy as np
from PIL import Image
from typing import Optional, List
from selenium import webdriver
import time 
from docx import Document
import fitz
import re
from transformers import BertTokenizerFast, BertForTokenClassification
import torch

import pandas as pd



def clean_indexes(df): 
    # Strip whitespace 
    df.columns = [col for col in df.columns.str.strip()]
    
    # Lowercasing 
    df.columns = [col for col in df.columns.str.lower()]
    
    # Append underscores 
    df.columns = [col.replace(" ", "_") if " " in col else col for col in df.columns]
    
    return df

def parse_dom_tags(tags): 
    d = dict()
    tags = [tag.replace("class=", "class_=") if "class=" in tag else tag for tag in tags.split(",")]
    
    for sep_tag in tags: 
        sep_tag = sep_tag.strip('[]').strip()
        key, val = sep_tag.split('=')
        val = val.strip("\"'")
        d[key] = val
        
    return d

driver1 = webdriver.Chrome()
        
def parse_javascript_body(url, dom_type, driver=driver1, **kwargs): 
    try: 
        driver.get(url)
    except Exception as e: 
        print(f'[ERROR] Failed to load URL, reason: {e}')
        return None 
    
    current_height = 0
    increment_scroll = 600
    end_height = driver.execute_script('return document.body.scrollHeight')
    found = False 
    
    driver_object = None
    
    while not found: 
        
        driver.execute_script(f'window.scrollTo(0, {current_height})')
        time.sleep(1)
        current_height += increment_scroll
        
        parser = BeautifulSoup(driver.page_source, "html.parser")
        body = parser.find(dom_type, **kwargs)
        
        if body: 
            found = True 
            print("Found item by scrolling...")
            driver_object = body
            break
        if current_height > end_height:
            print("No item found by scrolling...")
            break
        
    return driver_object


def multi_page_scrape(url, num, dom_header_type, **kwargs): 
    match = re.search(r"(\d+)(?=\D*$)", url)
    scraped_urls = []
    scraped_texts = []

    if not match:
        print(f"No number found in URL: {url}")
        return [], []
    
    if match: 
        page_number = match.group(1)
        print(url)
        secondary_urls, secondary_texts = complete_text_url_extraction(url, dom_header_type=dom_header_type, **kwargs)
        
        if secondary_urls: 
            scraped_urls.append(secondary_urls)
        if secondary_texts: 
            scraped_texts.append(secondary_texts)

        for _ in range(num):
            
            incremented = str(int(page_number) + 1)
            page_number = incremented
            new_url = url[:match.start()] + incremented + url[match.end():]
            print(new_url)
            secondary_urls, secondary_texts = complete_text_url_extraction(new_url, dom_header_type=dom_header_type, **kwargs)
            if secondary_urls:   
                scraped_urls.append(secondary_urls)
            if secondary_texts: 
                scraped_texts.append(secondary_texts)

    flattened_urls = [url for sublist in scraped_urls for url in sublist]
    flattened_texts = [text for text in scraped_texts]
    return flattened_urls, flattened_texts



def complete_text_url_extraction(url: str, dom_header_type: str, driver=driver1, **kwargs): 
    """Takes a URL, html_elements, and additional parameters like style, class, id and scrapes all the nested text 
    and URLs. Returns all the relevant texts and list of URLs. 

    Args:
        url (str): A URL of the website page 
        dom_header_type (str): HTML element type (e.g., div, p, header)

    Returns:
        _type_: _description_
    """
    
    html_response = requests.get(url)
    if html_response.ok: 
        parser = BeautifulSoup(html_response.content, 'html.parser')
        
    else:

        driver.get(url)
        
        if driver.page_source: 
            parser = BeautifulSoup(driver.page_source, 'html.parser')
            
        else: 
            print(f'{url} throws an error')
            return None, None
        
    header = parser.find(dom_header_type, **kwargs)
    
    if header is None:
            print(f'Trying to scrape via Selenium because requests was blocked.')
            header = parse_javascript_body(url, dom_header_type, driver=driver, **kwargs)
        
            if header is None: 
                print(f"Could not find {dom_header_type} with {kwargs} on {url}")
                
                return [], []
    
    # Extract all urls 
    sourced_links = header.find_all('a', href=True)
    print("Retrieving all URLs...")
    
    retrieved_urls = [link.get('href') for link in sourced_links]
    
    # Extract all texts
    
    print("Extracting all texts...")
    texts = header.get_text(separator='\n', strip=True)
    
    print("-----------------------")
    
    return retrieved_urls, texts

    




def dynamic_scraper(url, dom_header_type, multi_page=False, num=None, **kwargs):
    if multi_page: 
        print('Scraping multi-paged website')
        try: 
            scraped_urls, texts = multi_page_scrape(url, num, dom_header_type, **kwargs)
            return scraped_urls, texts
        except Exception as e:
            print(f"Error: {e}")
            return [], []
    else: 
        try: 
            scraped_urls, texts = complete_text_url_extraction(url, dom_header_type, **kwargs)
            return scraped_urls, texts   
        except Exception as e: 
            print(f'Error: {e}')
            return [], [] 


def ocr_extract_for_images(url):
    response = requests.get(url)
    if response.ok:
        img = Image.open(io.BytesIO(response.content))
        return pytesseract.image_to_string(img)
    else:
        print(f"Failed to retrieve image. Status code: {response.status_code}")
        return None
    

def text_extract_for_word_docs(url_docx):
    response = requests.get(url_docx)
    if response.ok:
        stream_doc = io.BytesIO(response.content)
        doc = Document(stream_doc)
        return '\n'.join([p.text for p in doc.paragraphs])
    else:
        print(f"Failed to retrieve Word document. Status code: {response.status_code}")
        return None

def text_extract_for_pdfs(url_pdfs):
    response = requests.get(url_pdfs)
    if response.ok:
        with io.BytesIO(response.content) as pdf_stream:
            doc = fitz.open(stream=pdf_stream.read(), filetype="pdf")
            if doc.page_count > 30:
                print("Document with more than 30 pages", url_pdfs)
                return None
            return '\n'.join(page.get_text() for page in doc)
    else:
        print(f"Failed to retrieve PDF. Status code: {response.status_code}")
        return None
        
# def ocr_extract_for_images(url_img, driver): 
    
#     valid_ext = ('.img', '.png', '.jpg', '.jpeg')
#     if url_img.endswith(valid_ext): 
        
#         response = driver.get(url_img)
        
#         if response: 
#             data = response.page_source 
#             streamed_image = Image.open(io.BytesIO(data))
            
#             return pytesseract.image_to_sttring(streamed_image)
        
#         else: 
            
#             return ""




# The first step is to handle non-HTML website items 
def url_validation(url_list: List[str], base_url: Optional[str] = None): 
    
    http_start = ('https://', 'http://')
    
    validated_url_list = []
    
    if not url_list: 
        return None 

    for url in url_list: 
        if not url: 
            continue 
        
        if url.startswith(http_start): 
            validated_url_list.append(url)
            continue 
        
        if base_url: 
            ## Try 1 
            parsed_url = urlparse(base_url)
            url_structure = f'{parsed_url.scheme}://{parsed_url.netloc}'
            joined_url = urljoin(url_structure, url)
            print(joined_url)
            
            try:
                response = requests.get(joined_url, timeout=2)
                if response.ok: 
                    validated_url_list.append(response.url) ## I put .url so that it can detect redirected pages
                    continue 
            except requests.RequestException:
                pass
            
            
            ## Try 2
            inherited_with_path = f"{parsed_url.scheme}://{parsed_url.netloc}{parsed_url.path}"
            joined_with_path = urljoin(inherited_with_path, url)
            try:
                response = requests.get(joined_with_path, timeout=2)
                if response.ok:
                    validated_url_list.append(response.url)
                    continue
            except requests.RequestException:
                pass

            print(f"Failed to validate URL: {url}")
        
        else: 
            pass 
        
    return validated_url_list if validated_url_list else None 





# 1. URL validation -> makes sure the link is a URL and it works 
# 2. handle downloadable files -> checks to see if PDF or image extensions -> if neither, returns the URL as a notdownloadable HTML link
# 3. PDF handler returns texts in a list
# 4. OCR image handler returns texts in a list 
# 5. Combine pipeline 




# def extract_text_from_pdfs(url: str): 
#     """Handles non-HTML extensions and extract all texts from the PDF. 

#     Args:
#         url (str): A URL of the website.

#     Returns:
#         _type_: _description_
#     """
#     ## validation
#     valid_ext_set = ('.pdf', '.txt')
    
#     if url.endswith(valid_ext_set): 
#         response = requests.get(url)
        
#         if response.status_code != 200: 
#             print(f"Error: Received status code {response.status_code}")
#             return ""
        
#         stream_data = response.content

#         # Stream the file data
#         # Resource: https://pymupdf.readthedocs.io/en/latest/how-to-open-a-file.html
#         doc = pymupdf.Document(stream=stream_data)
        
#         ## Exclude extraordinarily lengthy PDFs 
#         if doc.page_count > 20: 
#             pass 
#         # Resource: https://pymupdf.readthedocs.io/en/latest/app1.html 
#         return "\n".join(page.get_text("text", 
#                                        flags=pymupdf.TEXT_INHIBIT_SPACES & ~pymupdf.TEXT_PRESERVE_WHITESPACE) 
#                          for page in doc)
#     else: 
#         return ""
    


# def is_downloadable(url, base_url): 
#     validated_path = url_validation(url, base_url)
    
#     downloadable_files = ('.img', '.png', '.jpg', '.jpeg', '.txt', '.pdf', '.mp3', '.doc', 
#                           '.docx', '.csv', '.xlsx', '.ppt', '.pptx', '.zip')
#     if validated_path: 
#         if validated_path.endswith(downloadable_files): 
#             return True
#         else: 
#             return None
#     else: 
#         return None 
    
    


def handle_downloadable_websites(url): 
    converted_texts = None
    image_ext = ('.img', '.png', '.jpg', '.jpeg')
    file_ext = ('.txt', '.pdf')
    docs_ext = ('.docx',)
    
    if url is None: 
        return None 
    
    url_lower = url.lower()
    if url_lower.endswith(image_ext): 
        print('Extracting texts from images.')
        converted_texts = ocr_extract_for_images(url)
    elif url_lower.endswith(file_ext): 
        print('Extracting text from PDFs.')
        converted_texts = text_extract_for_pdfs(url)
    elif url_lower.endswith(docs_ext):
        print('Extracting text from Word documents.')
        converted_texts = text_extract_for_word_docs(url)
    else: 
        print('Not downloadable file.')
        
    return converted_texts



# def iterate_through_urls(url_list, base_url): 
    
#     not_downloadable = []
#     texts_list = []
#     for url in url_list: 
#         if is_downloadable(url, base_url): 
#             texts  = handle_downloadable_websites(url, base_url)
#             texts_list.append(texts)
#         if not is_downloadable(url, base_url): 
#             not_downloadable.append(url)
    
#     return texts_list, not_downloadable


def dedup_list(url_list):
    return list(set(url_list)) if isinstance(url_list, list) else []

def chain_func(*functions):
    def chained_function(x):
        result = x
        for func in functions:
            result = func(result)
        return result
    return chained_function

def remove_citations(text): 
    # 1) Remove references or bibliography section 
    text = re.split(r'\nReferences\n', text, flags=re.IGNORECASE)[0]
    text = re.split(r'\nBibliography\n', text, flags=re.IGNORECASE)[0]
    return text




## Data structure we want 
## URL || extracted texts and URLs || all the metadata || for each URL we can handle more non-HTML dataframe 

# def iterate_through_urls_with_pdfs(list_of_urls): 

#     not_downloadable = []
#     rows = [ ]
    
#     for url in list_of_urls: 
#         if is_downloadable(url): 

#             extracted_texts = extract_text_from_pdfs(url)
#             text_dct = {
#                 'topic': 'diet_supplements',
#                 'jurisdiction': 'MA',
#                 'polit_affil': None,
#                 'doc_type': 'pdf',
#                 'genre': 'academic',
#                 'polarity': None,
#                 'sentiment': None,
#                 'original_language': 1,
#                 'texts': extracted_texts
#             }
            
#             rows.append(text_dct)
            
#         else: 
#             not_downloadable.append(url)
            
#     return rows, not_downloadable 
    

# def data_entry(texts, row, **kwargs): 
#     """Takes the text and annotates it with the corresponding attributes"""
    
#     data_dictionary = {
#         "url": row['URL'],
#         "sub_url": row['URL_sub'],
#         "organization_name": row['Organization'],
#         "org_type": row['Type'],
#         "polit_aff": row['Political Affiliation'],
#         "year_established": row['Year Established'],
#         "region_state": row['Region of State'],
#         "state": row['State'],
#         "texts": texts,
#         "topic": None,
#         "text_modal": None,
#         "polarity": None,
#         "sentiment": None   
#     }
    
#     data_dictionary.update(kwargs)
    
#     return data_dictionary


model_name = "AventIQ-AI/named-entity-recognition-for-information-extraction"
ner_model = BertForTokenClassification.from_pretrained(model_name)
bert_tokenizer = BertTokenizerFast.from_pretrained(model_name)


def see_internal_comp(document, tokenizer=bert_tokenizer, chunk_size=512, model=ner_model):
    document = document.replace('\n', ' ')
    tokens = tokenizer(document, return_tensors='pt', truncation=False)
    input_ids = tokens['input_ids'][0]
    attention_mask = tokens['attention_mask'][0]
    
    # Split into chunks of chunk_size
    chunks = [(input_ids[i:i+chunk_size], attention_mask[i:i+chunk_size]) for i in range(0, len(input_ids), chunk_size)]
    
    final_texts = []
    
    for chunk_input_ids, chunk_attention_mask in chunks:
        # Add batch dimension
        chunk_input_ids = chunk_input_ids.unsqueeze(0)
        chunk_attention_mask = chunk_attention_mask.unsqueeze(0)
        
        with torch.no_grad():
            outputs = model(input_ids=chunk_input_ids, attention_mask=chunk_attention_mask)
        
        logits = outputs.logits
        preds = torch.argmax(logits, dim=2)[0].cpu().numpy()
        tokens_text = tokenizer.convert_ids_to_tokens(chunk_input_ids[0])
        
        for token, pred in zip(tokens_text, preds):
            print(token, pred)


def brute_remove_entities(document): 
    entities_list = ['Austin', 'Texas', 'TX', 'Dallas', 'San Antonio', 'Houston', 'El Paso', 'Arlington', 'Texan', 'Texans']
    
    for entity in entities_list: 
        document = document.replace(entity, '[PAD]')
    
    return document 
    
# def remove_entities_inference(document, tokenizer=bert_tokenizer, model=ner_model, chunk_size=512):
#     # Tokenize entire document
#     document = document.replace('\n', ' ')
#     tokens = tokenizer(document, return_tensors='pt', truncation=False)
#     input_ids = tokens['input_ids'][0]
#     attention_mask = tokens['attention_mask'][0]
    
#     # Split into chunks of chunk_size
#     chunks = [(input_ids[i:i+chunk_size], attention_mask[i:i+chunk_size]) for i in range(0, len(input_ids), chunk_size)]
    
#     final_texts = []
    
#     for chunk_input_ids, chunk_attention_mask in chunks:
#         # Add batch dimension
#         chunk_input_ids = chunk_input_ids.unsqueeze(0)
#         chunk_attention_mask = chunk_attention_mask.unsqueeze(0)
        
#         with torch.no_grad():
#             outputs = model(input_ids=chunk_input_ids, attention_mask=chunk_attention_mask)
        
#         logits = outputs.logits
#         preds = torch.argmax(logits, dim=2)[0].cpu().numpy()
#         tokens_text = tokenizer.convert_ids_to_tokens(chunk_input_ids[0])
    
#         out_tokens = []
        
        
#         for token, pred in zip(tokens_text, preds): 
#             if token in tokenizer.all_special_tokens:
#                 continue
#             if pred != 0:
#                 if out_tokens and out_tokens[-1] == '[PAD]': 
#                     continue
#                 out_tokens.append('[PAD]')
#             else: 
#                 if re.match(r"[.?!':;,\-\)]", token):
#                     out_tokens[-1] += token
#                     continue
#                 elif token.startswith('##'): 
#                     out_tokens[-1] += token[2:]
#                     continue
#                 out_tokens.append(token)
                
#         return " ".join(out_tokens)
            
            
def remove_entities_inference(document, tokenizer=bert_tokenizer, model=ner_model, chunk_size=512):
    
    # Tokenize entire document
    document = document.replace('\n', ' ')
    tokens = tokenizer(document, return_tensors='pt', truncation=False)
    input_ids = tokens['input_ids'][0]
    attention_mask = tokens['attention_mask'][0]
    
    # Split into chunks of chunk_size
    chunks = [(input_ids[i:i+chunk_size], attention_mask[i:i+chunk_size]) for i in range(0, len(input_ids), chunk_size)]
    
    final_texts = []
    
    for chunk_input_ids, chunk_attention_mask in chunks:
        # Add batch dimension
        chunk_input_ids = chunk_input_ids.unsqueeze(0)
        chunk_attention_mask = chunk_attention_mask.unsqueeze(0)
        
        with torch.no_grad():
            outputs = model(input_ids=chunk_input_ids, attention_mask=chunk_attention_mask)
        
        logits = outputs.logits
        preds = torch.argmax(logits, dim=2)[0].cpu().numpy()
        tokens_text = tokenizer.convert_ids_to_tokens(chunk_input_ids[0])
    
        out_tokens = []
        
        for token, pred in zip(tokens_text, preds): 
            if token in tokenizer.all_special_tokens:
                continue
            if pred != 0:
                if out_tokens and out_tokens[-1] == '[PAD]': 
                    continue
                out_tokens.append('[PAD]')
            else: 
                if re.match(r"[.?!':;,\-\)]", token):
                    if out_tokens:
                        out_tokens[-1] += token
                    else:
                        out_tokens.append(token)
                    continue
                elif token.startswith('##'): 
                    if out_tokens:
                        out_tokens[-1] += token[2:]
                    else:
                        out_tokens.append(token[2:])
                    continue
                out_tokens.append(token)
        
        final_texts.extend(out_tokens)  # Collect results for all chunks
    
    return " ".join(final_texts)


text= "The Texas Department of State Health Services is responsible for public health. One of TX DPH's job is to protect Texans. It is one of many state agencies that receives funding from the federal government."

remove_entities_inference(text)
    #         if token in tokenizer.all_special_tokens:
    #             continue
    #         # Replace 1 with the label ID for your target entity (e.g., ORG, PER, etc.)
    #         if pred != 0:  
    #             if token.startswith("##"):
    #                 final_texts[-1] += token[2:]
    #             else: 
    #                 final_texts.append(token)
    
    # return final_texts

        

# https://huggingface.co/AventIQ-AI/named-entity-recognition-for-information-extraction


    
    
    
    